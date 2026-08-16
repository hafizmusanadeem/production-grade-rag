from pathlib import Path

import logfire
from docx import Document as DocxDocument

from .base import BaseLoader, LoadedDocument


class DOCXLoader(BaseLoader):
    @logfire.instrument("Load DOCX {file_path=}", extract_args=("file_path",))
    def load(self, file_path: str | Path) -> list[LoadedDocument]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {path}")

        doc = DocxDocument(str(path))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        content = "\n\n".join(parts).strip()
        if not content:
            raise ValueError(f"No extractable text found in DOCX: {path}")

        metadata = self._base_metadata(path, "docx")
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)

        docs = [LoadedDocument(page_content=content, metadata=metadata)]
        self._record_load_metrics(
            path,
            docs,
            file_type="docx",
            paragraph_count=len(doc.paragraphs),
            table_count=len(doc.tables),
        )
        return docs
